# Module: Generated_word.py
from docx import Document

def generate_wordpad(filename, document_content):
    """
    Generate a clean Word document from structured or plain content.
    """
    # Create a new Word document
    doc = Document()

    # Add title or heading
    doc.add_heading("Legal Divorce Notice", level=1)

    if not document_content.strip():
        doc.add_paragraph("No content provided.")
    else:
        lines = document_content.strip().splitlines()

        for line in lines:
            line = line.strip()
            if not line:
                continue

            if line.startswith("# "):
                doc.add_heading(line[2:].strip(), level=1)
            elif line.startswith("## "):
                doc.add_heading(line[3:].strip(), level=2)
            elif line.startswith("- "):
                doc.add_paragraph(line[2:].strip(), style='List Bullet')
            elif line.startswith("**") and line.count("**") >= 2:
                parts = line.split("**")
                para = doc.add_paragraph()
                bold = False
                for p in parts:
                    run = para.add_run(p.strip())
                    run.bold = bold
                    bold = not bold
            else:
                doc.add_paragraph(line)

    # Save the file (should end in .docx)
    if not filename.endswith(".docx"):
        filename += ".docx"
    doc.save(filename)
